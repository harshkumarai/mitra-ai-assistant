"""Document text extractors for PDF, DOCX, and TXT formats."""

import os
from backend.utils.logger import get_logger

logger = get_logger(__name__)

# Lazy imports to prevent server crash if dependencies fail
_pypdf_available = True
try:
    import pypdf
except ImportError:
    _pypdf_available = False
    logger.warning("pypdf is not installed. PDF parsing will not be supported.")

_docx_available = True
try:
    import docx
except ImportError:
    _docx_available = False
    logger.warning("python-docx is not installed. DOCX parsing will not be supported.")


def parse_pdf(file_path: str) -> str:
    """Extract text from all pages of a PDF file.

    Args:
        file_path: Absolute path to the PDF file.

    Returns:
        Extracted text as a single string.
    """
    if not _pypdf_available:
        return "[Error: PDF parser is not installed on the server.]"
    
    text_content = []
    try:
        reader = pypdf.PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(f"--- Page {i+1} ---\n{page_text}")
        return "\n\n".join(text_content)
    except Exception as exc:
        logger.error("Failed to parse PDF '%s': %s", file_path, exc)
        return f"[Error parsing PDF: {exc}]"


def parse_docx(file_path: str) -> str:
    """Extract paragraphs and table text from a Microsoft Word DOCX file.

    Args:
        file_path: Absolute path to the DOCX file.

    Returns:
        Extracted text as a single string.
    """
    if not _docx_available:
        return "[Error: Word DOCX parser is not installed on the server.]"

    text_content = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Also parse tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
                    
        return "\n".join(text_content)
    except Exception as exc:
        logger.error("Failed to parse DOCX '%s': %s", file_path, exc)
        return f"[Error parsing DOCX: {exc}]"


def parse_txt(file_path: str) -> str:
    """Read a raw plain text file.

    Args:
        file_path: Absolute path to the text file.

    Returns:
        File contents as a string.
    """
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as exc:
        logger.error("Failed to read TXT file '%s': %s", file_path, exc)
        return f"[Error reading text file: {exc}]"


def parse_document(file_path: str) -> str:
    """Detect extension and parse document into a text representation.

    Supported formats: .pdf, .docx, .txt, .md, .csv

    Args:
        file_path: Path to the target document.

    Returns:
        The extracted document text context.
    """
    if not os.path.exists(file_path):
        return f"[Error: File '{file_path}' does not exist on the server.]"

    _, ext = os.path.splitext(file_path.lower())
    logger.info("Parsing file %s (extension: %s)", file_path, ext)

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in (".txt", ".md", ".csv", ".json", ".xml", ".html", ".css", ".js", ".py", ".ts"):
        return parse_txt(file_path)
    else:
        return f"[Unsupported file parsing extension: {ext}]"
